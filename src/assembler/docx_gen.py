from docx import Document
from docx.shared import Pt
from src.memory.schema import FullBook
import os

class DOCXGenerator:
    def __init__(self):
        self.output_dir = "output"
        os.makedirs(self.output_dir, exist_ok=True)

    def generate(self, book: FullBook) -> str:
        file_path = os.path.join(self.output_dir, f"{book.outline.title.replace(' ', '_')}.docx")
        doc = Document()

        # Front Matter
        for key, content in book.front_matter.items():
            doc.add_heading(key.replace('_', ' ').title(), level=1)
            doc.add_paragraph(content)
            doc.add_page_break()

        # Chapters
        for chapter in book.chapters:
            doc.add_heading(f"Chapter {chapter.chapter_number}: {chapter.title}", level=2)
            doc.add_paragraph(chapter.content)
            doc.add_page_break()

        # Back Matter
        for key, content in book.back_matter.items():
            doc.add_heading(key.replace('_', ' ').title(), level=2)
            doc.add_paragraph(content)
            doc.add_page_break()

        doc.save(file_path)
        print(f"[DOCXGenerator] DOCX created: {file_path}")
        return file_path
